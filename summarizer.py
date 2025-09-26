import os
import re
import math
from collections import Counter, defaultdict
from typing import List

from pypdf import PdfReader
from docx import Document
from typing import Optional
try:
    # OCR and PDF-to-image (optional, enabled if installed)
    from pdf2image import convert_from_path  # requires Poppler on Windows
except Exception:
    convert_from_path = None
try:
    import pypdfium2 as pdfium  # cross-platform PDF rendering, no Poppler needed
except Exception:
    pdfium = None
try:
    import pytesseract  # requires Tesseract OCR installed
except Exception:
    pytesseract = None

# Optional semantic embedding model
_EMB_MODEL = None
try:
    from sentence_transformers import SentenceTransformer, util as st_util  # optional
except Exception:
    SentenceTransformer = None
    st_util = None

# Optional fuzzy matching
try:
    from rapidfuzz import fuzz as rf_fuzz  # optional
except Exception:
    rf_fuzz = None

# Toggle OCR via environment variable (default: disabled)
USE_OCR = os.environ.get('USE_OCR', '0') == '1'
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

# Basic English stopwords list (compact, no external downloads)
STOPWORDS = set('''a about above after again against all am an and any are aren't as at be because been before being below
between both but by can't cannot could couldn't did didn't do does doesn't doing don't down during each few for from further
had hadn't has hasn't have haven't having he he'd he'll he's her here here's hers herself him himself his how how's i i'd i'll
i'm i've if in into is isn't it it's its itself let's me more most mustn't my myself no nor not of off on once only or other
ought our ours ourselves out over own same shan't she she'd she'll she's should shouldn't so some such than that that's the their
theirs them themselves then there there's these they they'd they'll they're they've this those through to too under until up very
was wasn't we we'd we'll we're we've were weren't what what's when when's where where's which while who who's whom why why's with
won't would wouldn't you you'd you'll you're you've your yours yourself yourselves'''.split())

SENTENCE_END_RE = re.compile(r"(?<=[.!?])\s+")
WORD_RE = re.compile(r"[A-Za-z][A-Za-z'\-]+")


def split_sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text.strip())
    # Split by sentence end regex, fallback if whole text becomes empty
    parts = SENTENCE_END_RE.split(text)
    sentences = [s.strip() for s in parts if s.strip()]
    return sentences or ([text] if text else [])


def tokenize(text: str) -> List[str]:
    return [w.lower() for w in WORD_RE.findall(text)]


def sentence_scores(sentences: List[str]):
    # Compute TF-IDF-like scoring with simple IDF approximation
    docs = [set(tokenize(s)) for s in sentences]
    df = Counter()
    for d in docs:
        for w in d:
            if w not in STOPWORDS:
                df[w] += 1
    N = len(sentences) or 1
    scores = []
    for s in sentences:
        tokens = tokenize(s)
        tf = Counter([w for w in tokens if w not in STOPWORDS])
        score = 0.0
        for w, f in tf.items():
            idf = math.log((N + 1) / (1 + df[w])) + 1.0
            score += (f / max(1, len(tokens))) * idf
        scores.append(score)
    return scores


def summarize_text(text: str, max_sentences: int = 5) -> str:
    sentences = split_sentences(text)
    if len(sentences) <= max_sentences:
        return " ".join(sentences)
    scores = sentence_scores(sentences)
    # Select top-k sentences, keep original order
    top_idx = sorted(sorted(range(len(sentences)), key=lambda i: scores[i], reverse=True)[:max_sentences])
    summary = " ".join([sentences[i] for i in top_idx])
    return summary

# --- Crisp, structured resume summarizer ---
_SKILL_ALIASES = {
    # Languages
    'python': ['python'],
    'java': ['java'],
    'javascript': ['javascript', 'js', 'node.js', 'nodejs', 'node'],
    'typescript': ['typescript', 'ts'],
    'c#': ['c#', '.net', 'dotnet'],
    'c++': ['c++'],
    'go': ['go', 'golang'],
    'sql': ['sql'],
    'nosql': ['nosql', 'mongo', 'mongodb', 'dynamodb', 'cassandra', 'redis'],
    # Frameworks / Libraries
    'react': ['react', 'react.js', 'reactjs'],
    'angular': ['angular', 'angular.js', 'angularjs'],
    'vue': ['vue', 'vue.js', 'vuejs'],
    'django': ['django'],
    'flask': ['flask'],
    'spring': ['spring', 'spring boot', 'springboot'],
    # Cloud / DevOps
    'aws': ['aws', 'amazon web services'],
    'azure': ['azure'],
    'gcp': ['gcp', 'google cloud'],
    'docker': ['docker'],
    'kubernetes': ['kubernetes', 'k8s'],
    'terraform': ['terraform'],
    'ci/cd': ['ci/cd', 'cicd', 'jenkins', 'github actions', 'gitlab ci'],
    # Data / ML
    'pandas': ['pandas'],
    'numpy': ['numpy'],
    'scikit-learn': ['scikit-learn', 'sklearn'],
    'ml': ['machine learning', 'ml'],
    'nlp': ['nlp', 'natural language processing'],
}

def _approx_years_experience(text: str) -> float:
    """Estimate total years of experience from explicit mentions and employment ranges."""
    import datetime
    t = text.lower()
    years = []
    # Explicit patterns like "X years", "X+ years"
    for m in re.finditer(r"(\d{1,2})\s*\+?\s*(?:\+|plus)?\s*(?:yrs|years|year)\b", t):
        try:
            years.append(float(m.group(1)))
        except Exception:
            pass
    # Ranges like 2018-2022, or "2019 to 2023"
    this_year = datetime.datetime.utcnow().year
    for m in re.finditer(r"(20\d{2}|19\d{2})\s*(?:-|to|–|—)\s*(present|current|20\d{2}|19\d{2})", t):
        start = int(m.group(1))
        end_raw = m.group(2)
        if end_raw in ('present', 'current'):
            end = this_year
        else:
            try:
                end = int(end_raw)
            except Exception:
                end = start
        if end >= start:
            years.append(max(0.0, float(end - start)))
    if not years:
        return 0.0
    # Heuristic: cap outrageous totals, take max as conservative signal
    est = max(years)
    return round(est, 1)


def _extract_strong_skills(text: str, top_n: int = 6) -> list:
    t = text.lower()
    counts = {}
    for canon, aliases in _SKILL_ALIASES.items():
        c = 0
        for a in aliases:
            # word boundary match to avoid false positives
            c += len(re.findall(rf"\b{re.escape(a)}\b", t))
        if c > 0:
            counts[canon] = c
    # Sort by frequency desc, then alphabetically
    ordered = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))
    return [k for k, _ in ordered[:top_n]]


def summarize_crisp(text: str, max_skills: int = 6) -> str:
    """
    Produce a concise, recruiter-friendly summary with years of experience and strong skills.
    Fallback to short extractive summary if we cannot detect signals.
    """
    yrs = _approx_years_experience(text)
    skills = _extract_strong_skills(text, top_n=max_skills)

    parts = []
    if yrs > 0:
        parts.append(f"Experience: ~{yrs} years")
    if skills:
        parts.append("Strong in: " + ", ".join(skills))

    if parts:
        return " | ".join(parts)
    # Fallback: extractive summary, short
    return summarize_text(text, max_sentences=3)

# --- Pointwise structured summary ---
_ROLE_KEYWORDS = [
    'software engineer', 'senior software', 'staff engineer', 'principal engineer', 'data scientist',
    'machine learning', 'ml engineer', 'full stack', 'frontend', 'backend', 'devops', 'site reliability', 'sre',
    'product manager', 'project manager', 'qa engineer', 'test engineer', 'security engineer', 'cloud engineer'
]

_DEGREE_PATTERNS = [
    r"\b(b\.?e\.?|btech|b\.tech|bachelor(?:'s)?|bs|bsc)\b",
    r"\b(m\.?e\.?|mtech|m\.tech|master(?:'s)?|ms|msc|mca|mba)\b",
    r"\b(phd|ph\.d\.|doctorate)\b"
]

def _find_roles(text: str, top_n: int = 3) -> list:
    t = text.lower()
    counts = {}
    for kw in _ROLE_KEYWORDS:
        c = len(re.findall(rf"\b{re.escape(kw)}\b", t))
        if c:
            counts[kw] = c
    ordered = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))
    return [k for k, _ in ordered[:top_n]]


def _find_degrees(text: str, top_n: int = 2) -> list:
    t = text.lower()
    found = []
    for pat in _DEGREE_PATTERNS:
        if re.search(pat, t):
            found.append(re.sub(r"\\\\b|\\\\.", '', pat))
    return found[:top_n]


def _find_highlights(text: str, top_n: int = 4) -> list:
    # Pick sentences with action verbs and metrics
    sentences = split_sentences(text)
    action_verbs = r"(led|built|designed|implemented|developed|optimized|improved|reduced|increased|migrated|automated)"
    metric = r"(\b\d+\s*%|\b\d+\s*(?:k|m|million|billion)|\$\s*\d+)"
    candidates = []
    for s in sentences:
        sc = 0
        if re.search(action_verbs, s, flags=re.I):
            sc += 1
        if re.search(metric, s, flags=re.I):
            sc += 1
        if sc > 0:
            candidates.append((sc, len(s), s.strip()))
    # Sort by score desc, then shorter is better
    ordered = sorted(candidates, key=lambda x: (-x[0], x[1]))
    return [s for _, __, s in ordered[:top_n]]


def summarize_pointwise(text: str) -> str:
    """
    Return a concise bullet list covering the main recruiter signals.
    Bullets: Experience, Strong Skills, Roles, Education, Highlights (up to 4)
    """
    yrs = _approx_years_experience(text)
    skills = _extract_strong_skills(text, top_n=8)
    roles = _find_roles(text)
    degrees = _find_degrees(text)
    highlights = _find_highlights(text)

    bullets = []
    if yrs > 0:
        bullets.append(f"- Experience: ~{yrs} years")
    if skills:
        bullets.append("- Strong skills: " + ", ".join(skills))
    if roles:
        bullets.append("- Roles: " + ", ".join(roles))
    if degrees:
        bullets.append("- Education: " + ", ".join(degrees))
    if highlights:
        bullets.append("- Highlights:")
        for h in highlights:
            bullets.append(f"  • {h}")

    if not bullets:
        return summarize_text(text, max_sentences=3)
    return "\n".join(bullets)


def _extract_companies(text: str, top_n: int = 4) -> list:
    t = text
    # Patterns: "at Company", "@ Company", uppercase styled names, suffixes like Inc, LLC, Ltd, Technologies, Labs
    company_patterns = [
        r"(?:at|@)\s+([A-Z][A-Za-z0-9&\-\. ]{2,})",
        r"\b([A-Z][A-Za-z0-9&\-\. ]+\s+(?:Inc\.|Inc|LLC|Ltd\.|Ltd|Technologies|Labs|Systems|Solutions))\b",
    ]
    candidates = []
    for pat in company_patterns:
        for m in re.finditer(pat, t):
            cand = m.group(1).strip().rstrip('.,;:')
            if len(cand.split()) <= 6:
                candidates.append(cand)
    # Count frequency
    counts = Counter([c for c in candidates if c])
    ordered = [c for c, _ in counts.most_common()]
    # Deduplicate preserving order
    seen = set()
    uniq = []
    for c in ordered:
        k = c.lower()
        if k not in seen:
            seen.add(k)
            uniq.append(c)
    return uniq[:top_n]


_CERT_PATTERNS = [
    r"aws\s+certified[\w\s\-]*",
    r"azure\s+certified[\w\s\-]*",
    r"gcp\s+certified[\w\s\-]*",
    r"pmp",
    r"scrum\s*master",
    r"cissp",
    r"compTIA\s*(?:security\+|network\+)",
]

def _find_certifications(text: str, top_n: int = 3) -> list:
    t = text.lower()
    found = []
    for pat in _CERT_PATTERNS:
        for m in re.finditer(pat, t):
            found.append(m.group(0))
    # Normalize and dedupe
    norm = []
    seen = set()
    for f in found:
        v = re.sub(r"\s+", " ", f.strip())
        if v not in seen:
            seen.add(v)
            norm.append(v)
    return norm[:top_n]


def summarize_pointwise_structured(text: str) -> tuple:
    """Return (text_bullets, html_list) for pointwise summary with companies and qualifications."""
    yrs = _approx_years_experience(text)
    skills = _extract_strong_skills(text, top_n=8)
    roles = _find_roles(text)
    degrees = _find_degrees(text)
    highlights = _find_highlights(text)
    companies = _extract_companies(text)
    certs = _find_certifications(text)

    bullets = []
    if yrs > 0:
        bullets.append(f"- Experience: ~{yrs} years")
    if companies:
        bullets.append("- Companies: " + ", ".join(companies))
    if skills:
        bullets.append("- Strong skills: " + ", ".join(skills))
    if roles:
        bullets.append("- Roles: " + ", ".join(roles))
    quals = []
    if degrees:
        quals.append("Degrees: " + ", ".join(degrees))
    if certs:
        quals.append("Certifications: " + ", ".join(certs))
    if quals:
        bullets.append("- Qualifications: " + "; ".join(quals))
    if highlights:
        bullets.append("- Highlights:")
        for h in highlights:
            bullets.append(f"  • {h}")

    if not bullets:
        txt = summarize_text(text, max_sentences=3)
        html = f"<p>{re.sub(r'&', '&amp;', txt)}</p>"
        return txt, html

    # Build HTML list from bullets
    def esc(s: str) -> str:
        return (s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
    html_lines = [
        '<ul class="ats-summary">'
    ]
    for b in bullets:
        if b.startswith("- Highlights:"):
            html_lines.append('<li><strong>Highlights:</strong><ul>')
        elif b.startswith("  • "):
            html_lines.append(f"<li>{esc(b[4:])}</li>")
        else:
            # Split label and value
            if ': ' in b:
                label, value = b[2:].split(': ', 1)
                html_lines.append(f"<li><strong>{esc(label)}:</strong> {esc(value)}</li>")
            else:
                html_lines.append(f"<li>{esc(b[2:])}</li>")
    # Close nested list if any highlights were added
    if any(b.startswith("  • ") for b in bullets):
        html_lines.append('</ul></li>')
    html_lines.append('</ul>')
    html = "".join(html_lines)
    txt = "\n".join(bullets)
    return txt, html


# -------------------- ATS Scoring --------------------
def _tf(text: str) -> Counter:
    tokens = [w for w in tokenize(text) if w not in STOPWORDS]
    return Counter(tokens)


def _cosine_sim(c1: Counter, c2: Counter) -> float:
    if not c1 or not c2:
        return 0.0
    # Build union of keys
    keys = set(c1.keys()) | set(c2.keys())
    if not keys:
        return 0.0
    v1 = []
    v2 = []
    for k in keys:
        v1.append(c1.get(k, 0.0))
        v2.append(c2.get(k, 0.0))
    # Compute cosine
    import math as _m
    dot = sum(a*b for a, b in zip(v1, v2))
    n1 = _m.sqrt(sum(a*a for a in v1))
    n2 = _m.sqrt(sum(b*b for b in v2))
    if n1 == 0 or n2 == 0:
        return 0.0
    return float(dot / (n1 * n2))


def _keyword_overlap(resume_text: str, jd_text: str) -> float:
    """Compute overlap of strong skills and roles as a ratio [0,1]."""
    rs_skills = set(_extract_strong_skills(resume_text, top_n=12))
    jd_skills = set(_extract_strong_skills(jd_text, top_n=12))
    rs_roles = set(_find_roles(resume_text))
    jd_roles = set(_find_roles(jd_text))
    # Jaccard overlaps
    def jacc(a: set, b: set) -> float:
        if not a and not b:
            return 0.0
        inter = len(a & b)
        union = len(a | b) or 1
        return inter / union
    skill_score = jacc(rs_skills, jd_skills)
    role_score = jacc(rs_roles, jd_roles)
    # Weighted average favors skills
    return 0.7 * skill_score + 0.3 * role_score


def _fuzzy_has(hay: str, needle: str, threshold: int = 85) -> bool:
    """Return True if fuzzy ratio >= threshold. Requires rapidfuzz; else False."""
    if rf_fuzz is None:
        return False
    try:
        return int(rf_fuzz.partial_ratio(hay.lower(), needle.lower())) >= threshold
    except Exception:
        return False


def _extract_skills(text: str) -> set:
    t = text.lower()
    found = set()
    for canon, aliases in _SKILL_ALIASES.items():
        for a in aliases:
            # exact
            if re.search(rf"\b{re.escape(a)}\b", t):
                found.add(canon)
                break
            # fuzzy fallback (for variations like "ReactJS", "Typescript", spacing issues, etc.)
            if _fuzzy_has(t, a, threshold=88):
                found.add(canon)
                break
    return found


def _parse_jd_requirements(jd_text: str) -> tuple[set, set]:
    """Heuristically split JD skills into must-have vs nice-to-have."""
    t = jd_text.lower()
    must_markers = [
        r"must\s+have",
        r"required",
        r"minimum",
        r"at\s+least",
        r"strong\s+experience",
    ]
    nice_markers = [
        r"nice\s+to\s+have",
        r"preferred",
        r"plus",
        r"bonus",
        r"good\s+to\s+have",
    ]
    jd_skills = _extract_skills(jd_text)
    must = set()
    nice = set()
    for s in jd_skills:
        aliases = _SKILL_ALIASES.get(s, [s])
        # If any alias appears near a must marker, treat as must
        is_must = any(re.search(rf"({m}).{{0,50}}\b{re.escape(aliases[0])}\b|\b{re.escape(aliases[0])}\b.{{0,50}}({m})", t) for m in must_markers)
        is_nice = any(re.search(rf"({m}).{{0,50}}\b{re.escape(aliases[0])}\b|\b{re.escape(aliases[0])}\b.{{0,50}}({m})", t) for m in nice_markers)
        if is_must and not is_nice:
            must.add(s)
        elif is_nice and not is_must:
            nice.add(s)
    # Anything not classified defaults to must to be conservative
    unclassified = jd_skills - must - nice
    must |= unclassified
    return must, nice


def _split_resume_sections(text: str) -> dict:
    """Split resume into sections by common headings for section-weighted relevance."""
    t = text
    # Simple heading-based split
    headings = [
        'experience', 'work experience', 'professional experience', 'projects',
        'skills', 'technical skills', 'education', 'certifications', 'summary'
    ]
    indices = []
    for h in headings:
        for m in re.finditer(rf"\n\s*{re.escape(h)}\s*:?\s*\n", t, flags=re.I):
            indices.append((m.start(), h.lower()))
    indices.sort()
    sections = {}
    if not indices:
        sections['full'] = t
        return sections
    for i, (start, name) in enumerate(indices):
        end = indices[i+1][0] if i+1 < len(indices) else len(t)
        sections[name] = t[start:end]
    return sections


def _bm25_score(resume_text: str, jd_text: str) -> float:
    """Compute BM25 relevance of resume w.r.t JD (query=JD tokens). Return 0..1 normalized."""
    # Tokenize
    q_tokens = [w for w in tokenize(jd_text) if w not in STOPWORDS]
    d_tokens = [w for w in tokenize(resume_text) if w not in STOPWORDS]
    if not q_tokens or not d_tokens:
        return 0.0
    # Document frequencies approximated with query terms inside document
    tf = Counter(d_tokens)
    # IDF with low-N smoothing (treat single doc)
    # BM25 params
    k1 = 1.5
    b = 0.75
    avgdl = max(1.0, len(d_tokens))
    dl = len(d_tokens)
    # Build query term set (dedupe)
    q_unique = list(dict.fromkeys(q_tokens))
    score = 0.0
    for term in q_unique:
        f = tf.get(term, 0)
        if f == 0:
            continue
        # pseudo-IDF
        idf = math.log(1 + (1 / (1 + 0)))  # ~ ln(2)
        num = f * (k1 + 1)
        den = f + k1 * (1 - b + b * dl / avgdl)
        score += idf * (num / den)
    # Normalize with a sigmoid to 0..1
    return 1.0 / (1.0 + math.exp(-score / 10.0))


def _embedding_similarity(text1: str, text2: str) -> float:
    """Return semantic similarity (0..1) using SentenceTransformers if available; else 0.0."""
    global _EMB_MODEL
    if SentenceTransformer is None or st_util is None:
        return 0.0
    try:
        if _EMB_MODEL is None:
            # Small, fast, widely available
            _EMB_MODEL = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
        e1 = _EMB_MODEL.encode(text1[:5000], convert_to_tensor=True, normalize_embeddings=True)
        e2 = _EMB_MODEL.encode(text2[:5000], convert_to_tensor=True, normalize_embeddings=True)
        sim = float(st_util.cos_sim(e1, e2).cpu().numpy().flatten()[0])
        # Map from [-1,1] to [0,1] if needed, but cos_sim already in [-1,1]; clamp
        return max(0.0, min(1.0, (sim + 1.0) / 2.0))
    except Exception:
        return 0.0


def _infer_title_tokens(text: str) -> set:
    """Heuristic: infer role/title tokens from text (first line and role keywords)."""
    tokens = set()
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if lines:
        first = lines[0].lower()
        for w in tokenize(first):
            if len(w) > 2 and w not in STOPWORDS:
                tokens.add(w)
    # Add role keywords
    for r in _find_roles(text):
        for w in r.split():
            if len(w) > 2:
                tokens.add(w)
    return tokens


def compute_ats_score(resume_text: str, jd_text: str, job_title: str | None = None) -> dict:
    """Return a dict with normalized ATS score [0,100] and detailed components used by modern screeners.

    Components (all 0..1 before aggregation):
      - bm25: lexical relevance of resume to JD (BM25)
      - must_match: fraction of JD must-have skills present in resume
      - nice_match: fraction of JD nice-to-have skills present in resume
      - years_alignment: proximity of years in resume vs requirement
      - section_alignment: weighted BM25 across Experience/Skills/Education sections
      - cosine, keyword: kept for reference/back-compat
    """
    resume_text = (resume_text or '').strip()
    jd_text = (jd_text or '').strip()
    if not resume_text or not jd_text:
        return {"score": 0, "cosine": 0.0, "keyword": 0.0, "years_alignment": 0.0,
                "bm25": 0.0, "embed": 0.0, "title_match": 0.0,
                "must_match": 0.0, "nice_match": 0.0, "section_alignment": 0.0,
                "matched_skills": [], "missing_must_skills": [], "missing_nice_skills": [], "capped": False}

    # Cosine on TF vectors
    c1 = _tf(resume_text)
    c2 = _tf(jd_text)
    cosine = _cosine_sim(c1, c2)

    # Keyword overlap on skills/roles
    keyword = _keyword_overlap(resume_text, jd_text)

    # Years alignment (if JD mentions years required)
    import re as _re
    yrs_resume = _approx_years_experience(resume_text)
    yrs_req = 0.0
    for m in _re.finditer(r"(\d{1,2})\s*\+?\s*(?:yrs|years|year)\s*(?:of)?\s*(?:experience|exp)\b", jd_text.lower()):
        try:
            yrs_req = max(yrs_req, float(m.group(1)))
        except Exception:
            pass
    years_alignment = 0.0
    if yrs_req > 0:
        # Smooth proximity: 1.0 if resume >= req; else decay as ratio
        years_alignment = min(1.0, (yrs_resume / yrs_req) if yrs_req else 0.0)
    else:
        # If no explicit requirement, be neutral
        years_alignment = 0.5 if yrs_resume > 0 else 0.3

    # BM25 relevance
    bm25 = _bm25_score(resume_text, jd_text)

    # Embedding similarity (semantic)
    embed = _embedding_similarity(resume_text, jd_text)

    # Skills matching (must/nice)
    jd_must, jd_nice = _parse_jd_requirements(jd_text)
    rs_skills = _extract_skills(resume_text)
    # Fuzzy match JD skills against resume raw text as a secondary signal
    matched_must = []
    missing_must = []
    for s in sorted(jd_must):
        if s in rs_skills:
            matched_must.append(s)
        else:
            aliases = _SKILL_ALIASES.get(s, [s])
            if any(_fuzzy_has(resume_text, a, threshold=88) for a in aliases):
                matched_must.append(s)
            else:
                missing_must.append(s)
    matched_nice = []
    missing_nice = []
    for s in sorted(jd_nice):
        if s in rs_skills:
            matched_nice.append(s)
        else:
            aliases = _SKILL_ALIASES.get(s, [s])
            if any(_fuzzy_has(resume_text, a, threshold=88) for a in aliases):
                matched_nice.append(s)
            else:
                missing_nice.append(s)
    must_match = (len(matched_must) / len(jd_must)) if jd_must else 0.5  # neutral if none identified
    nice_match = (len(matched_nice) / len(jd_nice)) if jd_nice else 0.5

    # Section-weighted alignment
    sections = _split_resume_sections(resume_text)
    weights = {
        'experience': 0.5, 'work experience': 0.5, 'professional experience': 0.5, 'projects': 0.4,
        'skills': 0.3, 'technical skills': 0.3, 'education': 0.2, 'certifications': 0.2, 'summary': 0.2,
        'full': 0.4,
    }
    sec_score = 0.0
    total_w = 0.0
    for name, txt in sections.items():
        w = weights.get(name, 0.1)
        sec_score += w * _bm25_score(txt, jd_text)
        total_w += w
    section_alignment = (sec_score / total_w) if total_w else 0.0

    # Title alignment
    title_match = 0.0
    try:
        if job_title:
            rs_title = _infer_title_tokens(resume_text)
            jd_title = _infer_title_tokens(job_title)
        else:
            rs_title = _infer_title_tokens(resume_text)
            jd_title = _infer_title_tokens(jd_text)
        inter = len(rs_title & jd_title)
        union = len(rs_title | jd_title) or 1
        title_match = inter / union
    except Exception:
        title_match = 0.0

    # Weighted aggregation (tuned heuristically). Increase weight on must-match and embeddings.
    # Retuned weights: emphasize must-have coverage and semantic/lexical relevance
    raw = (
        0.20 * bm25 +
        0.22 * embed +
        0.36 * must_match +
        0.06 * nice_match +
        0.06 * years_alignment +
        0.06 * section_alignment +
        0.03 * title_match +
        0.008 * cosine +
        0.006 * keyword
    )

    # Cap score if missing must-haves
    capped = False
    if missing_must:
        capped = True
        # Cap scales with how many must-haves are missing (stricter if many missing)
        miss_ratio = len(missing_must) / max(1, len(jd_must))
        cap = 0.55 - 0.15 * min(1.0, miss_ratio)  # between ~0.40 and 0.55
        raw = min(raw, cap)
    # Normalize to 0-100
    score = int(round(max(0.0, min(1.0, raw)) * 100))
    return {
        "score": score,
        "cosine": round(float(cosine), 4),
        "keyword": round(float(keyword), 4),
        "years_alignment": round(float(years_alignment), 4),
        "bm25": round(float(bm25), 4),
        "embed": round(float(embed), 4),
        "title_match": round(float(title_match), 4),
        "must_match": round(float(must_match), 4),
        "nice_match": round(float(nice_match), 4),
        "section_alignment": round(float(section_alignment), 4),
        "matched_skills": matched_must + matched_nice,
        "missing_must_skills": missing_must,
        "missing_nice_skills": missing_nice,
        "capped": capped,
    }


def summarize_for_jd(text: str) -> dict:
    """Produce both text and HTML summaries using the structured pointwise summarizer."""
    txt, html = summarize_pointwise_structured(text)
    return {"text": txt, "html": html}


# -------------------- Keyword-only ATS --------------------
def _extract_keywords(text: str, min_len: int = 3, top_k: int | None = None) -> set:
    """Extract a set of keywords from free text using tokenization and stopword filtering.
    Optionally limit to top_k by frequency. Also merges detected skills for better coverage.
    """
    toks = [t for t in tokenize(text) if t not in STOPWORDS and len(t) >= min_len]
    freq = Counter(toks)
    words = list(freq.keys())
    if top_k is not None and top_k > 0:
        words = [w for w, _ in freq.most_common(top_k)]
    # Merge with detected skills for domain relevance
    skills = _extract_skills(text)
    return set(words) | set(skills)


def compute_ats_keywords(resume_text: str, jd_text: str, top_k: int | None = None, cap_per_term: int = 3) -> dict:
    """Compute a keyword-only ATS score using frequency-weighted precision/recall/F1.

    Returns dict with fields:
      - score: 0..100 (F1 * 100)
      - precision, recall, f1: 0..1 metrics (frequency-weighted with per-term cap)
      - matched_keywords: list of overlapped keywords
      - missing_keywords: list of JD keywords not found in resume
      - jd_keywords, resume_keywords: for transparency
    """
    resume_text = (resume_text or '').strip()
    jd_text = (jd_text or '').strip()
    if not resume_text or not jd_text:
        return {
            "score": 0, "precision": 0.0, "recall": 0.0, "f1": 0.0,
            "matched_keywords": [], "missing_keywords": [],
            "jd_keywords": [], "resume_keywords": []
        }

    jd_kw = _extract_keywords(jd_text, top_k=top_k)
    rs_kw = _extract_keywords(resume_text, top_k=top_k)
    if not jd_kw or not rs_kw:
        return {
            "score": 0, "precision": 0.0, "recall": 0.0, "f1": 0.0,
            "matched_keywords": [], "missing_keywords": list(sorted(jd_kw)),
            "jd_keywords": list(sorted(jd_kw)), "resume_keywords": list(sorted(rs_kw))
        }

    # Build frequency dicts with stopword filtering and cap per term to avoid gaming
    jd_tokens = [t for t in tokenize(jd_text) if t not in STOPWORDS and len(t) >= 3]
    rs_tokens = [t for t in tokenize(resume_text) if t not in STOPWORDS and len(t) >= 3]
    jd_freq = Counter(jd_tokens)
    rs_freq = Counter(rs_tokens)

    # Only count terms that are in extracted keyword sets
    jd_terms = {k: min(c, cap_per_term) for k, c in jd_freq.items() if k in jd_kw}
    rs_terms = {k: min(c, cap_per_term) for k, c in rs_freq.items() if k in rs_kw}

    # Also boost known skills slightly by adding +1 (capped) if present in skills set
    jd_skills = _extract_skills(jd_text)
    rs_skills = _extract_skills(resume_text)
    for s in jd_skills:
        # Map skill to its canonical token key (use the skill name itself)
        jd_terms[s] = min(jd_terms.get(s, 0) + 1, cap_per_term)
    for s in rs_skills:
        rs_terms[s] = min(rs_terms.get(s, 0) + 1, cap_per_term)

    # Compute frequency-weighted overlap
    all_terms = set(jd_terms.keys()) | set(rs_terms.keys())
    tp = 0
    jd_total = 0
    rs_total = 0
    matched = []
    for term in all_terms:
        jf = jd_terms.get(term, 0)
        rf = rs_terms.get(term, 0)
        jd_total += jf
        rs_total += rf
        common = min(jf, rf)
        if common > 0:
            matched.append(term)
        tp += common

    precision = (tp / rs_total) if rs_total > 0 else 0.0
    recall = (tp / jd_total) if jd_total > 0 else 0.0
    f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) > 0 else 0.0
    score = int(round(f1 * 100))

    return {
        "score": score,
        "precision": round(float(precision), 4),
        "recall": round(float(recall), 4),
        "f1": round(float(f1), 4),
        "matched_keywords": list(sorted(set(matched))),
        "missing_keywords": list(sorted(jd_kw - rs_kw)),
        "jd_keywords": list(sorted(jd_kw)),
        "resume_keywords": list(sorted(rs_kw)),
    }


def _detect_tesseract_cmd() -> str:
    if pytesseract is None:
        return ''
    # 1) Respect environment variable first
    tess_cmd = os.environ.get('TESSERACT_CMD')
    if tess_cmd and os.path.exists(tess_cmd):
        return tess_cmd
    # 2) Try common Windows install location
    common = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
    if os.path.exists(common):
        return common
    # 3) Try Program Files (x86)
    common86 = r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe"
    if os.path.exists(common86):
        return common86
    return ''


def _detect_poppler_path() -> str:
    # 1) Respect environment variable first
    pp = os.environ.get('POPPLER_PATH', '')
    if pp and os.path.isdir(pp):
        return pp
    # 2) Try a couple of common locations
    guesses = [
        r"C:\\poppler\\Library\\bin",
        r"C:\\Program Files\\poppler\\Library\\bin",
    ]
    for g in guesses:
        if os.path.isdir(g):
            return g
    return ''


def _render_pdf_with_pdfium(path: str):
    """Render PDF pages to PIL Images using pypdfium2 if available."""
    if pdfium is None:
        return []
    try:
        pdf = pdfium.PdfDocument(path)
        images = []
        # Render at higher scale for better OCR
        for i in range(len(pdf)):
            page = pdf[i]
            pil_image = page.render(scale=2).to_pil()
            images.append(pil_image)
        return images
    except Exception:
        return []


def extract_text_from_pdf(path: str) -> str:
    # 1) Try PyPDF2 first
    pieces = []
    try:
        with open(path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                try:
                    pieces.append(page.extract_text() or '')
                except Exception:
                    continue
    except Exception:
        pieces = []

    text = "\n".join(pieces).strip()
    # 2) Fallback to pdfminer if PyPDF2 gave little or no text
    if (not text or len(text) < 20) and pdfminer_extract_text is not None:
        try:
            text = (pdfminer_extract_text(path) or '').strip()
        except Exception:
            pass

    # If still empty/very short, try OCR fallback (only if enabled)
    if USE_OCR and (not text or len(text) < 20) and pytesseract is not None:
        try:
            # Prefer pdfium rendering if available
            images = _render_pdf_with_pdfium(path)
            if not images and convert_from_path is not None:
                poppler_path = _detect_poppler_path() or None  # e.g., C:\\poppler-xx\\Library\\bin
                images = convert_from_path(path, dpi=300, poppler_path=poppler_path)
            ocr_chunks = []
            # Configure tesseract path if provided
            tess_cmd = _detect_tesseract_cmd()  # e.g., C:\\Program Files\\Tesseract-OCR\\tesseract.exe
            if tess_cmd and pytesseract is not None:
                pytesseract.pytesseract.tesseract_cmd = tess_cmd
            for img in images:
                try:
                    ocr_chunks.append(pytesseract.image_to_string(img) or '')
                except Exception:
                    continue
            text = "\n".join(ocr_chunks).strip()
        except Exception:
            pass

    # Normalize whitespace
    text = re.sub(r"\s+", " ", text)
    return text


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def extract_text_from_image(path: str) -> str:
    """Extract text from a raster image using OCR if available."""
    try:
        from PIL import Image
        if pytesseract is None:
            return ''
        tess_cmd = _detect_tesseract_cmd()
        if tess_cmd:
            pytesseract.pytesseract.tesseract_cmd = tess_cmd
        img = Image.open(path)
        text = pytesseract.image_to_string(img) or ''
        return re.sub(r"\s+", " ", text.strip())
    except Exception:
        return ''


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(path)
    if ext in ('.docx', '.doc'):
        return extract_text_from_docx(path)
    if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp'):
        return extract_text_from_image(path)
    # Fallback: treat as plain text
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            return re.sub(r"\s+", " ", content)
    except Exception:
        return ''
